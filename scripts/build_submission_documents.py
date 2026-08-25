#!/usr/bin/env python3
"""Build clean, editable cover-letter and highlights DOCX files."""

from __future__ import annotations

import argparse
from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
BLACK = "000000"
MUTED = "666666"


def set_font(run, size: float = 11, *, bold: bool = False, italic: bool = False, color: str = BLACK) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08
    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.45)
    bullet.paragraph_format.first_line_indent = Inches(-0.22)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.12


def add_inline(paragraph, text: str) -> None:
    token = re.compile(r"(\*\*.*?\*\*|\*[^*]+?\*)")
    position = 0
    for match in token.finditer(text):
        if match.start() > position:
            set_font(paragraph.add_run(text[position:match.start()]))
        value = match.group(0)
        if value.startswith("**"):
            set_font(paragraph.add_run(value[2:-2]), bold=True)
        else:
            set_font(paragraph.add_run(value[1:-1]), italic=True)
        position = match.end()
    if position < len(text):
        set_font(paragraph.add_run(text[position:]))


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size=18, bold=True, color=BLUE)


def build_cover(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure(doc)
    add_title(doc, "Cover Letter")
    for raw in lines[1:]:
        line = raw.rstrip()
        if not line:
            continue
        p = doc.add_paragraph()
        if line in {"Editor-in-Chief  ", "Sincerely,"}:
            p.paragraph_format.keep_with_next = True
        add_inline(p, line.replace("  ", ""))
    core = doc.core_properties
    core.title = "Cover Letter — Beyond Average Latency"
    core.subject = "Submission to Journal of Systems Architecture"
    core.author = "Akhtar Saleem"
    core.last_modified_by = "Akhtar Saleem"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_highlights(source: Path, output: Path) -> None:
    bullets = [line[2:].strip() for line in source.read_text(encoding="utf-8").splitlines() if line.startswith("- ")]
    if len(bullets) != 5:
        raise ValueError(f"Expected exactly five highlights, found {len(bullets)}")
    if any(len(value) > 85 for value in bullets):
        raise ValueError("A highlight exceeds Elsevier's 85-character guidance")
    doc = Document()
    configure(doc)
    add_title(doc, "Highlights")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.add_run("Beyond Average Latency"), size=11, italic=True, color=MUTED)
    for value in bullets:
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(value), size=11)
    core = doc.core_properties
    core.title = "Highlights — Beyond Average Latency"
    core.subject = "Research highlights for journal submission"
    core.author = "Akhtar Saleem"
    core.last_modified_by = "Akhtar Saleem"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--paper-dir", type=Path, default=Path("paper"))
    args = parser.parse_args()
    build_cover(args.paper_dir / "cover_letter_draft.md", args.paper_dir / "cover_letter.docx")
    build_highlights(args.paper_dir / "highlights.md", args.paper_dir / "highlights.docx")
    print(args.paper_dir / "cover_letter.docx")
    print(args.paper_dir / "highlights.docx")


if __name__ == "__main__":
    main()
