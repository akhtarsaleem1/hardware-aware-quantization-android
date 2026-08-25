#!/usr/bin/env python3
"""Build a styled, accessible journal DOCX from the finalized Markdown source."""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "666666"
LIGHT_FILL = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "000000"
PORTRAIT_WIDTH_DXA = 9360
LANDSCAPE_WIDTH_DXA = 12960


@dataclass
class Token:
    kind: str
    value: object
    level: int = 0


def set_font(run, name: str = "Calibri", size: float | None = None, *, bold=None, italic=None, color: str | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, *, before: float = 0, after: float = 8, line: float = 1.333, keep_with_next: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_font(run, size=9, color=MUTED)


def configure_section(section, *, landscape: bool = False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.page_width = Inches(11 if landscape else 8.5)
    section.page_height = Inches(8.5 if landscape else 11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    header.is_linked_to_previous = True
    if not header.paragraphs[0].text:
        p = header.paragraphs[0]
        p.text = "Quantization rankings on Android | Original research manuscript"
        set_spacing(p, after=0, line=1.0)
        set_font(p.runs[0], size=8.5, color=MUTED)
    footer = section.footer
    footer.is_linked_to_previous = True
    has_page_field = any(
        (node.text or "").strip().startswith("PAGE")
        for node in footer.paragraphs[0]._p.xpath(".//w:instrText")
    )
    if not has_page_field:
        add_page_number(footer.paragraphs[0])


def patch_style(style, *, size: float, color: str = BLACK, bold: bool = False, italic: bool = False, before: float = 0, after: float = 8, line: float = 1.333) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def configure_styles(doc: Document) -> None:
    patch_style(doc.styles["Normal"], size=11, after=8, line=1.333)
    patch_style(doc.styles["Title"], size=20, color=BLACK, bold=True, after=8, line=1.0)
    patch_style(doc.styles["Subtitle"], size=11, color=MUTED, after=4, line=1.0)
    patch_style(doc.styles["Heading 1"], size=16, color=BLUE, bold=True, before=18, after=10, line=1.0)
    patch_style(doc.styles["Heading 2"], size=13, color=BLUE, bold=True, before=12, after=6, line=1.0)
    patch_style(doc.styles["Heading 3"], size=12, color=DARK_BLUE, bold=True, before=8, after=4, line=1.0)
    patch_style(doc.styles["Caption"], size=9, color=BLACK, italic=False, before=4, after=4, line=1.0)
    for name in ("List Bullet", "List Number"):
        patch_style(doc.styles[name], size=11, after=4, line=1.208)
        doc.styles[name].paragraph_format.left_indent = Inches(0.375)
        doc.styles[name].paragraph_format.first_line_indent = Inches(-0.194)


def citation_text(text: str) -> str:
    def repl(match: re.Match[str]) -> str:
        keys = re.findall(r"P(\d+)", match.group(1))
        return "[" + ", ".join(str(int(value)) for value in keys) + "]"
    return re.sub(r"\[@([^\]]+)\]", repl, text)


def add_inline(paragraph, text: str) -> None:
    text = citation_text(text)
    token = re.compile(r"(\*\*.*?\*\*|`.*?`|\*[^*]+?\*|\[[^\]]+\]\([^\)]+\))")
    position = 0
    for match in token.finditer(text):
        if match.start() > position:
            set_font(paragraph.add_run(text[position:match.start()]))
        value = match.group(0)
        if value.startswith("**"):
            run = paragraph.add_run(value[2:-2])
            set_font(run, bold=True)
        elif value.startswith("`"):
            run = paragraph.add_run(value[1:-1])
            set_font(run, name="Consolas", size=9.5)
        elif value.startswith("*"):
            run = paragraph.add_run(value[1:-1])
            set_font(run, italic=True)
        else:
            label = value[1:value.index("](")]
            set_font(paragraph.add_run(label), color=BLUE)
        position = match.end()
    if position < len(text):
        set_font(paragraph.add_run(text[position:]))


def strip_front_matter(text: str) -> tuple[dict[str, str], str]:
    metadata: dict[str, str] = {}
    if not text.startswith("---\n"):
        return metadata, text
    _, front, body = text.split("---", 2)
    title = re.search(r'^title:\s*"(.*)"\s*$', front, flags=re.M)
    journal = re.search(r'^journal:\s*"(.*)"\s*$', front, flags=re.M)
    author = re.search(r'^\s*-\s*name:\s*"(.*)"\s*$', front, flags=re.M)
    affiliation = re.search(r'^\s*affiliation:\s*"(.*)"\s*$', front, flags=re.M)
    email = re.search(r'^\s*email:\s*"(.*)"\s*$', front, flags=re.M)
    orcid = re.search(r'^\s*orcid:\s*"(.*)"\s*$', front, flags=re.M)
    metadata["title"] = title.group(1) if title else "Untitled manuscript"
    metadata["journal"] = journal.group(1) if journal else ""
    metadata["author"] = author.group(1) if author else "[AUTHOR NAME]"
    metadata["affiliation"] = (
        affiliation.group(1) if affiliation else "[AFFILIATION]"
    )
    metadata["email"] = (
        email.group(1) if email else "[CORRESPONDING AUTHOR EMAIL]"
    )
    metadata["orcid"] = orcid.group(1) if orcid else ""
    return metadata, body.lstrip()


def tokenize(markdown: str) -> list[Token]:
    lines = markdown.splitlines()
    tokens: list[Token] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            tokens.append(Token("paragraph", " ".join(value.strip() for value in paragraph)))
            paragraph.clear()

    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            flush()
            index += 1
            continue
        image = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line.strip())
        if image:
            flush()
            tokens.append(Token("image", (image.group(1), image.group(2))))
            index += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            flush()
            tokens.append(Token("heading", heading.group(2), len(heading.group(1))))
            index += 1
            continue
        if line.lstrip().startswith("|"):
            flush()
            table_lines: list[str] = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = [[cell.strip() for cell in value.strip("|").split("|")] for value in table_lines]
            if len(rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
                rows.pop(1)
            tokens.append(Token("table", rows))
            continue
        bullet = re.match(r"^[-*]\s+(.*)$", line)
        number = re.match(r"^\d+\.\s+(.*)$", line)
        if bullet or number:
            flush()
            tokens.append(Token("list_bullet" if bullet else "list_number", (bullet or number).group(1)))
            index += 1
            continue
        paragraph.append(line)
        index += 1
    flush()
    return tokens


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], total: int) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def table_widths(rows: list[list[str]], total: int) -> list[int]:
    count = len(rows[0])
    weights = []
    for column in range(count):
        longest = max(len(row[column]) if column < len(row) else 0 for row in rows)
        weights.append(max(6, min(longest, 36)))
    scale = total / sum(weights)
    widths = [max(500, round(value * scale)) for value in weights]
    widths[-1] += total - sum(widths)
    return widths


def add_table(doc: Document, rows: list[list[str]], caption: str | None) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=True)
    if caption:
        p = doc.add_paragraph(style="Caption")
        add_inline(p, caption)
        p.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = table_widths(rows, LANDSCAPE_WIDTH_DXA)
    set_table_geometry(table, widths, LANDSCAPE_WIDTH_DXA)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if column_index in (0, len(values) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            set_spacing(p, after=0, line=1.0)
            add_inline(p, value)
            for run in p.runs:
                set_font(run, size=7.2, bold=row_index == 0, color=WHITE if row_index == 0 else BLACK)
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), DARK_BLUE)
                cell._tc.get_or_add_tcPr().append(shading)
        if row_index == 0:
            table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=False)


def add_image(doc: Document, source: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(source), width=Inches(6.4))
    drawing = run._r.find(qn("w:drawing"))
    if drawing is not None:
        doc_pr = drawing.find(".//" + qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", caption)
    p.paragraph_format.keep_with_next = True
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(cap, caption)


def bib_entries(text: str) -> list[dict[str, str]]:
    entries: list[dict[str, str]] = []
    position = 0
    while True:
        start = text.find("@", position)
        if start < 0:
            break
        brace = text.find("{", start)
        comma = text.find(",", brace)
        if brace < 0 or comma < 0:
            break
        depth = 1
        index = brace + 1
        while index < len(text) and depth:
            depth += (text[index] == "{") - (text[index] == "}")
            index += 1
        body = text[comma + 1:index - 1]
        entry = {"key": text[brace + 1:comma].strip()}
        field_pattern = re.compile(r"(?m)^\s*(\w+)\s*=\s*\{")
        for match in field_pattern.finditer(body):
            value_start = match.end()
            depth = 1
            cursor = value_start
            while cursor < len(body) and depth:
                depth += (body[cursor] == "{") - (body[cursor] == "}")
                cursor += 1
            entry[match.group(1).lower()] = body[value_start:cursor - 1].replace("{{", "").replace("}}", "").strip()
        entries.append(entry)
        position = index
    return sorted(entries, key=lambda value: int(re.sub(r"\D", "", value["key"])))


def add_references(doc: Document, entries: list[dict[str, str]]) -> None:
    for number, entry in enumerate(entries, start=1):
        authors = entry.get("author", "Unknown author").replace(" and ", ", ")
        venue = entry.get("journal") or entry.get("booktitle") or entry.get("howpublished") or ""
        identifier = f" doi:{entry['doi']}" if entry.get("doi") else ""
        if not identifier and entry.get("eprint"):
            identifier = f" arXiv:{entry['eprint']}"
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        set_spacing(p, after=5, line=1.0)
        title = entry.get("title", "").replace("{", "").replace("}", "")
        add_inline(p, f"{number}. {authors} ({entry.get('year', 'n.d.')}). {title}. {venue}.{identifier}")
        for run in p.runs:
            set_font(run, size=9)


def build(markdown: Path, bibliography: Path, output: Path) -> None:
    metadata, body = strip_front_matter(markdown.read_text(encoding="utf-8"))
    tokens = tokenize(body)
    references = bib_entries(bibliography.read_text(encoding="utf-8"))
    if len(references) < 25:
        raise ValueError("Bibliography has fewer than 25 parsed entries")
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])

    kicker = doc.add_paragraph()
    set_spacing(kicker, after=6, line=1.0)
    run = kicker.add_run("ORIGINAL RESEARCH MANUSCRIPT")
    set_font(run, size=9, bold=True, color=BLUE)
    title = doc.add_paragraph(style="Title")
    add_inline(title, metadata["title"])
    subtitle = doc.add_paragraph(style="Subtitle")
    add_inline(subtitle, f"{metadata['author']} | {metadata['affiliation']} | {metadata['email']}")
    if metadata["orcid"]:
        orcid_line = doc.add_paragraph(style="Subtitle")
        add_inline(orcid_line, f"ORCID: {metadata['orcid']}")
    journal = doc.add_paragraph(style="Subtitle")
    add_inline(journal, metadata["journal"])
    set_spacing(journal, after=16, line=1.0)

    pending_caption: str | None = None
    index = 0
    while index < len(tokens):
        token = tokens[index]
        if token.kind == "heading" and token.value == "References":
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, "References")
            add_references(doc, references)
            break
        if token.kind == "paragraph" and str(token.value).startswith("**Table ") and index + 1 < len(tokens) and tokens[index + 1].kind == "table":
            pending_caption = str(token.value).strip("*")
            index += 1
            continue
        if token.kind == "heading":
            style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[token.level]
            p = doc.add_paragraph(style=style)
            add_inline(p, str(token.value))
            p.paragraph_format.keep_with_next = True
        elif token.kind == "paragraph":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_spacing(p)
            add_inline(p, str(token.value))
        elif token.kind in ("list_bullet", "list_number"):
            p = doc.add_paragraph(style="List Bullet" if token.kind == "list_bullet" else "List Number")
            add_inline(p, str(token.value))
            p.paragraph_format.keep_together = True
        elif token.kind == "table":
            add_table(doc, token.value, pending_caption)
            pending_caption = None
        elif token.kind == "image":
            caption, relative = token.value
            source = (markdown.parent / relative).resolve()
            if not source.is_file():
                raise FileNotFoundError(source)
            add_image(doc, source, caption)
        index += 1

    core = doc.core_properties
    core.title = metadata["title"]
    core.subject = "Mobile AI quantization repeatability research manuscript"
    core.author = metadata["author"]
    core.last_modified_by = metadata["author"]
    core.keywords = "Android, LiteRT, quantization, mobile inference, repeatability"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--markdown", type=Path, default=Path("paper/manuscript_main.md"))
    parser.add_argument("--bibliography", type=Path, default=Path("paper/references.bib"))
    parser.add_argument("--output", type=Path, default=Path("paper/manuscript_main.docx"))
    args = parser.parse_args()
    build(args.markdown, args.bibliography, args.output)
    print(args.output)


if __name__ == "__main__":
    main()
